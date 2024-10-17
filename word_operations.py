from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from voice_engine import speak, takeCommand


def search_for_document(file=""):
    while file == "" or file == "none":
        speak("Please tell the name of the file")
        file = takeCommand().lower()
    folders = os.listdir()
    if file + ".docx" in folders:
        speak("File is present in this directory")
        return 1
    else:
        speak("File is Not Present in this directory")
        return 0


def making_new_document():
    speak("Please give the name of the document")
    file_name = takeCommand().lower()

    if search_for_document(file=file_name) == 1:
        speak("Do you want to open existing file or want to make another file with another name")
        ans = takeCommand().lower()
        if "open existing file" in ans:
            opening_existing_document(file_name)
        else:
            making_new_document()
    else:
        document = Document()
        document.save(file_name + ".docx")
        speak(f"New document {file_name}.docx has been created")


def opening_existing_document(file=""):
    commandlist = ['read the document', "add a paragraph",
                   "add a table", "add an image", "add a heading", "add a list"]
    if search_for_document(file) == 1:
        speak("What do you want to do now, 1. read the document, 2. add a paragraph, 3. add a table, 4. add an image, 5. add a heading, 6. add a List")
        query = takeCommand().lower()
        if query in commandlist:
            if "add a paragraph" in query:
                add_paragraph(file)
            elif "add a heading" in query:
                add_heading(file)
            elif "add a list" in query:
                add_list(file)
            elif "read the document" in query:
                read_the_document(file)
        else:
            speak("Command not recognized. Please try again.")


def add_paragraph(file):
    document = Document(file + ".docx")
    paragraph = document.add_paragraph()
    speak("Please say what you want to enter in the paragraph. Say 'quit para' when you're done.")

    while True:
        content = takeCommand().lower()
        if "quit para" in content:
            break
        paragraph.add_run(content.capitalize() + ". ")

    document.save(file + ".docx")
    speak("Paragraph has been added successfully")


def add_heading(file):
    document = Document(file + ".docx")
    speak("Please say the heading you want to add")
    heading = takeCommand()
    speak("What should be the level of the heading? (0-9)")
    level = int(takeCommand())
    document.add_heading(heading, level=level)
    document.save(file + ".docx")
    speak("Heading has been added successfully")


def add_list(file):
    document = Document(file + ".docx")
    speak("Please mention the style for this list. Your options are: 1. List Number, 2. List Bullet")
    style = "List Number" if "number" in takeCommand().lower() else "List Bullet"

    speak("Now speak the items of the list. Say 'quit list' when you're done.")
    while True:
        item = takeCommand()
        if "quit list" in item.lower():
            break
        document.add_paragraph(item, style=style)

    document.save(file + '.docx')
    speak("List has been added successfully")


def read_the_document(file):
    document = Document(file + ".docx")
    for para in document.paragraphs:
        speak(para.text)
